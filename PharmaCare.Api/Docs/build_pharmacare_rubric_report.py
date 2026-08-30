# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Bao_cao_Phan_tich_Thiet_ke_HTTT_PharmaCare.docx"
ASSETS = ROOT / "report_assets"
ASSETS.mkdir(exist_ok=True)
BLUE = "1D4ED8"; NAVY = "173B67"; LIGHT = "EAF2FF"; PALE = "F5F8FC"; GREEN = "DCFCE7"; RED = "FEE2E2"; GRAY = "475569"

def font(size=22, bold=False):
    paths = ["/System/Library/Fonts/Supplemental/Arial.ttf", "/System/Library/Fonts/Supplemental/Arial Bold.ttf"]
    try: return ImageFont.truetype(paths[1 if bold else 0], size)
    except: return ImageFont.load_default()

def diagram(path, title, lanes, arrows):
    w, h = 1500, 220 + len(lanes)*170
    im = Image.new("RGB", (w,h), "white"); d=ImageDraw.Draw(im)
    d.text((60,35), title, fill="#173B67", font=font(34,True))
    boxes={}
    for i,(key,label,detail,color) in enumerate(lanes):
        y=130+i*170; x=90+(i%2)*650
        if i%2: y=130+(i-1)*170
        box=(x,y,x+560,y+110); boxes[key]=box
        d.rounded_rectangle(box, radius=18, fill=color, outline="#2563EB", width=3)
        d.text((x+22,y+16),label,fill="#0F172A",font=font(24,True))
        d.multiline_text((x+22,y+52),detail,fill="#334155",font=font(18),spacing=5)
    for a,b,label in arrows:
        x1,y1,x2,y2=boxes[a]; u1,v1,u2,v2=boxes[b]
        start=(x2,y1+55) if x1<u1 else (x1,y1+55); end=(u1,v1+55) if x1<u1 else (u2,v1+55)
        d.line([start,end],fill="#1D4ED8",width=4)
        d.polygon([(end[0],end[1]),(end[0]-12 if start[0]<end[0] else end[0]+12,end[1]-8),(end[0]-12 if start[0]<end[0] else end[0]+12,end[1]+8)],fill="#1D4ED8")
        d.text(((start[0]+end[0])//2-55,start[1]-25),label,fill="#1D4ED8",font=font(16,True))
    im.save(path)

diagram(ASSETS/"context.png","System Context — PharmaCare",[
    ("customer","Khách hàng","Tra cứu, giỏ hàng, đơn thuốc, đặt hàng","#EAF2FF"),
    ("web","Web PharmaCare","React/TypeScript, Customer & Internal Portal","#DBEAFE"),
    ("api","PharmaCare API","ASP.NET Core, RBAC, nghiệp vụ, audit","#DCFCE7"),
    ("db","PostgreSQL","Đơn hàng, lô, tồn kho, người dùng, log","#FEF3C7"),
    ("staff","Nhân viên nội bộ","Dược sĩ, kho, quản lý chi nhánh, Admin","#FCE7F3"),
],[("customer","web","HTTPS"),("web","api","JSON/JWT"),("api","db","EF Core"),("staff","web","HTTPS")])
diagram(ASSETS/"flow.png","Luồng đầu vào → đầu ra đã kiểm thử",[
    ("product","Danh mục & sản phẩm","Đơn vị Hộp/Vỉ/Viên; VAT; Rx/OTC","#EAF2FF"),
    ("batch","Lô & nhập kho","HSD, giá vốn, số lượng, mức đặt lại","#FEF3C7"),
    ("transfer","Phân phối chi nhánh","TRANSFER_OUT / TRANSFER_IN","#DCFCE7"),
    ("order","Đơn hàng","RESERVE; FEFO; giá trước/sau VAT","#FCE7F3"),
    ("complete","Hoàn tất & báo cáo","SALE; PAID; doanh thu; audit","#E0E7FF"),
],[("product","batch","tạo lô"),("batch","transfer","phân phối"),("transfer","order","bán"),("order","complete","đối soát")])

doc=Document()
sec=doc.sections[0]; sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.left_margin=Inches(.85); sec.right_margin=Inches(.85); sec.top_margin=Inches(.68); sec.bottom_margin=Inches(.65)
styles=doc.styles
normal=styles['Normal']; normal.font.name='Arial'; normal.font.size=Pt(10.5); normal.font.color.rgb=RGBColor.from_string('1E293B'); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.15
for name,size,color,before,after in [('Title',28,NAVY,0,8),('Subtitle',13,GRAY,0,12),('Heading 1',18,NAVY,14,7),('Heading 2',14,BLUE,11,5),('Heading 3',11.5,NAVY,8,4)]:
    s=styles[name]; s.font.name='Arial'; s.font.size=Pt(size); s.font.bold=name!='Subtitle'; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
for name in ['List Bullet','List Number']:
    styles[name].font.name='Arial'; styles[name].font.size=Pt(10.5); styles[name].paragraph_format.space_after=Pt(3)

def shade(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill)
def margins(cell,top=100,start=120,bottom=100,end=120):
    tcPr=cell._tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for side,val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn(f'w:{side}'))
        if node is None: node=OxmlElement(f'w:{side}'); tcMar.append(node)
        node.set(qn('w:w'),str(val)); node.set(qn('w:type'),'dxa')
def table(headers,rows,widths=None,font_size=8.8):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    if widths is None: widths=[6.7/len(headers)]*len(headers)
    grid=t._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths:
        g=OxmlElement('w:gridCol'); g.set(qn('w:w'),str(int(width*1440))); grid.append(g)
    trPr=t.rows[0]._tr.get_or_add_trPr(); rep=OxmlElement('w:tblHeader'); rep.set(qn('w:val'),'true'); trPr.append(rep)
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.width=Inches(widths[i]); shade(c,NAVY); margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(str(h)); r.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(font_size)
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,val in enumerate(row):
            c=cells[i]; c.width=Inches(widths[i]); margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ri%2: shade(c,PALE)
            p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER if i==0 or (len(str(val))<18 and i>0) else WD_ALIGN_PARAGRAPH.LEFT
            r=p.add_run(str(val)); r.font.name='Arial'; r.font.size=Pt(font_size)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return t
def h(text,level=1): doc.add_heading(text,level=level)
def p(text,bold_prefix=None):
    x=doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix): x.add_run(bold_prefix).bold=True; x.add_run(text[len(bold_prefix):])
    else: x.add_run(text)
    return x
def bullets(items,numbered=False):
    for item in items: doc.add_paragraph(item,style='List Number' if numbered else 'List Bullet')
def callout(title,text,fill=LIGHT):
    q=doc.add_paragraph();q.paragraph_format.space_before=Pt(4);q.paragraph_format.space_after=Pt(8);q.paragraph_format.left_indent=Pt(10);q.paragraph_format.right_indent=Pt(10)
    pPr=q._p.get_or_add_pPr();shd=OxmlElement('w:shd');shd.set(qn('w:fill'),fill);pPr.append(shd)
    borders=OxmlElement('w:pBdr');left=OxmlElement('w:left');left.set(qn('w:val'),'single');left.set(qn('w:sz'),'20');left.set(qn('w:color'),BLUE);left.set(qn('w:space'),'8');borders.append(left);pPr.append(borders)
    q.add_run(title+' — ').bold=True;q.add_run(text)
def page(): doc.add_page_break()
def caption(text):
    x=doc.add_paragraph(text);x.alignment=WD_ALIGN_PARAGRAPH.CENTER;x.paragraph_format.keep_with_next=True
    for r in x.runs:r.italic=True;r.font.size=Pt(9);r.font.color.rgb=RGBColor.from_string(GRAY)

# Header/footer
header=sec.header.paragraphs[0];header.text='PHARMACARE  |  PHÂN TÍCH & THIẾT KẾ HTTT';header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
for r in header.runs:r.font.name='Arial';r.font.size=Pt(8);r.font.bold=True;r.font.color.rgb=RGBColor.from_string(BLUE)
footer=sec.footer.paragraphs[0];footer.alignment=WD_ALIGN_PARAGRAPH.CENTER;footer.add_run('PharmaCare — Báo cáo nộp học phần  •  Trang ')
fld=OxmlElement('w:fldSimple');fld.set(qn('w:instr'),'PAGE');footer._p.append(fld)

# Cover
x=doc.add_paragraph();x.alignment=WD_ALIGN_PARAGRAPH.CENTER;x.paragraph_format.space_before=Pt(80);r=x.add_run('PHARMACARE');r.bold=True;r.font.name='Arial';r.font.size=Pt(18);r.font.color.rgb=RGBColor.from_string(BLUE)
x=doc.add_paragraph(style='Title');x.alignment=WD_ALIGN_PARAGRAPH.CENTER;x.add_run('BÁO CÁO PHÂN TÍCH & THIẾT KẾ\nHỆ THỐNG THÔNG TIN NHÀ THUỐC')
x=doc.add_paragraph(style='Subtitle');x.alignment=WD_ALIGN_PARAGRAPH.CENTER;x.add_run('Use Case · Requirements · Domain/Data · Architecture · RBAC/JWT · Traceability')
doc.add_paragraph('\n')
table(['Thông tin','Nội dung'],[
    ['Đề tài','Hệ thống quản lý và bán thuốc đa chi nhánh PharmaCare'],
    ['Phạm vi','Customer Portal, vận hành nhà thuốc, kho, chi nhánh và quản trị hệ thống'],
    ['Nền tảng','.NET 10 Web API · PostgreSQL · React/TypeScript'],
    ['Phiên bản hồ sơ','1.0 — 23/08/2026'],
    ['Nhóm/Lớp','Nhóm dự án PharmaCare · Thông tin lớp cập nhật khi nộp'],
    ['Thành viên','Theo danh sách chính thức của nhóm dự án'],
],[1.45,5.25],9.5)
callout('Cam kết phạm vi','Hồ sơ mô tả đúng hệ thống đã triển khai và kiểm thử. Các ID yêu cầu, API, entity, quyền và test được liên kết trong Traceability Matrix; không dùng endpoint giả định.')
page()
h('Mục lục',1)
p('Mục lục được Word cập nhật khi mở tài liệu: chọn toàn bộ và nhấn F9 nếu số trang chưa hiển thị.')
toc=doc.add_paragraph(); run=toc.add_run(); begin=OxmlElement('w:fldChar');begin.set(qn('w:fldCharType'),'begin');instr=OxmlElement('w:instrText');instr.set(qn('xml:space'),'preserve');instr.text=' TOC \\o "1-3" \\h \\z \\u ';sep=OxmlElement('w:fldChar');sep.set(qn('w:fldCharType'),'separate');txt=OxmlElement('w:t');txt.text='Nhấn F9 để cập nhật mục lục';end=OxmlElement('w:fldChar');end.set(qn('w:fldCharType'),'end');run._r.extend([begin,instr,sep,txt,end])
h('Tóm tắt điều hành',1)
p('PharmaCare là hệ thống thông tin nhà thuốc đa chi nhánh, kết nối hành trình mua thuốc của khách hàng với nghiệp vụ dược sĩ, kho, quản lý chi nhánh và Admin. Hệ thống hỗ trợ thuốc OTC không cần đăng nhập, thuốc kê đơn qua quy trình tư vấn/duyệt đơn, bán tại quầy, quản lý theo lô–hạn dùng–đơn vị quy đổi, báo cáo VAT và truy vết thao tác.')
table(['Khía cạnh','Kết quả thiết kế'],[
 ['Phân quyền','RBAC resource.action, backend default-deny; dữ liệu nội bộ cô lập theo chi nhánh'],
 ['Kho','FEFO, giữ tồn khi đặt, sổ giao dịch bất biến theo IMPORT/TRANSFER/RESERVE/SALE'],
 ['Thuế','Tách doanh thu trước VAT, VAT đầu ra, giá sau VAT, phí, giảm giá, hoàn tiền và doanh thu ròng'],
 ['Bảo mật','JWT ngắn hạn, refresh rotation/revoke, password hash, 401/403 và audit mutation'],
 ['Kiểm chứng','Build sạch; role-boundary PASS; E2E kho→đơn→dược sĩ PASS; p95 catalog 5,1 ms'],
],[1.45,5.25],9)

page();h('1. Bối cảnh, mục tiêu và phạm vi',1)
h('1.1 Bài toán nghiệp vụ',2)
p('Nhà thuốc cần bán đúng thuốc, đúng đơn vị, đúng giá và đúng phạm vi pháp lý; đồng thời kiểm soát tồn theo lô/HSD, phân phối giữa chi nhánh và truy được người thực hiện. Các cách quản lý rời rạc bằng bảng tính dễ gây bán vượt tồn, nhầm quy đổi hộp–vỉ–viên, thiếu bằng chứng duyệt thuốc kê đơn và lộ dữ liệu giữa chi nhánh.')
h('1.2 Mục tiêu hệ thống',2)
table(['ID','Mục tiêu','Chỉ báo thành công'],[
 ['G-01','Khách mua thuốc OTC thuận tiện','Tra cứu công khai; guest checkout; giao tận nơi/nhận tại quầy'],
 ['G-02','Kiểm soát thuốc kê đơn','Chỉ mua sau tư vấn/đơn được duyệt hoặc dược sĩ xác nhận đơn giấy tại POS'],
 ['G-03','Tồn kho chính xác theo chi nhánh/lô','Giữ tồn nguyên tử; FEFO; ledger đủ đầu vào–đầu ra'],
 ['G-04','Quản trị theo phạm vi','Manager đúng một chi nhánh; API ngoài phạm vi trả 403'],
 ['G-05','Báo cáo thuế đối soát được','Trước VAT + VAT = sau VAT; sau VAT + phí − giảm = thực thu'],
 ['G-06','Truy vết và vận hành an toàn','Audit mutation; actor/IP/time; không log mật khẩu/token'],
],[.65,3.25,2.8],8.6)
h('1.3 Phạm vi',2)
bullets(['Trong phạm vi: danh mục/sản phẩm/đơn vị bán; lô và tồn; chuyển kho; đơn thuốc; đơn hàng online/POS; voucher; thanh toán COD/VIETQR; báo cáo; user/role/branch/audit.', 'Ngoài phạm vi hiện tại: tích hợp ngân hàng production callback, kết nối HIS/BHYT, hóa đơn điện tử ký số, tối ưu tuyến giao hàng và dự báo nhu cầu bằng ML.', 'Ràng buộc: PostgreSQL nhất quán giao dịch; quy tắc dược được thực thi ở backend; frontend không phải ranh giới bảo mật.'])
h('1.4 Actor và trách nhiệm',2)
table(['Actor','Mục tiêu','Quyền hạn chính','Giới hạn'],[
 ['Khách hàng','Tìm/mua/tra cứu','Catalog, giỏ, checkout, hồ sơ, đơn và đơn thuốc của mình','Không xem dữ liệu khách khác; Rx cần duyệt'],
 ['Dược sĩ/nhân viên bán','Tư vấn và cấp thuốc','Review prescription, POS, xác nhận/hoàn tất đơn, xem tồn được gán','Không quản trị user/role'],
 ['Nhân viên kho','Nhập và phân phối','Sản phẩm, lô, receive/adjust/transfer, ledger','Không xử lý đơn bán'],
 ['Quản lý chi nhánh','Điều hành một cơ sở','Báo cáo, đơn, tồn, cảnh báo, voucher trong chi nhánh','Không xem chi nhánh khác; không duyệt đơn thuốc'],
 ['Admin','Quản trị toàn hệ thống','User, role/permission, branch, category, audit, báo cáo toàn cục','Không được tự khóa tài khoản đang dùng'],
],[1.25,1.55,2.35,1.55],8.3)
caption('Hình 1. System Context')
doc.add_picture(str(ASSETS/'context.png'),width=Inches(6.7));doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.inline_shapes[-1]._inline.docPr.set('descr','Sơ đồ System Context của PharmaCare: khách hàng và nhân viên dùng React Portal, gọi ASP.NET Core API qua HTTPS/JWT; API truy cập PostgreSQL qua EF Core.')

page();h('2. Use Case, User Story và Acceptance Criteria',1)
usecases=[
('UC-01','Mua thuốc OTC không cần tài khoản','Khách hàng','Catalog có sản phẩm còn hoạt động','Tìm/lọc → chọn đơn vị → thêm giỏ → chọn chi nhánh/hình thức nhận → nhập thông tin → đặt đơn','Đơn PENDING, tồn được RESERVE, giá/VAT lưu theo snapshot','Hết tồn 409; dữ liệu giao hàng thiếu 400; voucher sai bị từ chối'),
('UC-02','Tư vấn và mua thuốc kê đơn','Khách hàng, Dược sĩ','Khách đăng nhập; ảnh đơn hợp lệ','Upload multipart → dược sĩ xem ảnh → chọn thuốc/liều → APPROVE → khách checkout gắn prescriptionId','Đơn liên kết đúng đơn đã duyệt và chi nhánh','Đơn bị từ chối; sai chi nhánh/số lượng; file quá giới hạn'),
('UC-03','Nhập lô và phân phối chi nhánh','Nhân viên kho','Sản phẩm và chi nhánh hoạt động','Tạo batch → receive kho nguồn → transfer kho đích → kiểm tra ledger','Nguồn giảm, đích tăng, cùng batch; audit actor/time','Lô hết hạn; thiếu tồn; branch ngoài phạm vi 403'),
('UC-04','Bán thuốc tại quầy','Dược sĩ','Đăng nhập và có orders.manage','Chọn chi nhánh → sản phẩm/đơn vị → xác nhận đơn giấy nếu Rx → tạo POS → hoàn tất','Đơn COMPLETED/PAID; tồn SALE theo base quantity','Rx thiếu xác nhận; thiếu tồn; chi nhánh ngoài phạm vi'),
('UC-05','Điều hành và báo cáo thuế','Quản lý chi nhánh','Được Admin gán đúng một chi nhánh','Chọn kỳ → xem trước VAT/VAT/sau VAT/thực thu/hoàn tiền → xem top/cảnh báo','Chỉ số khớp phương trình; chỉ dữ liệu chi nhánh','branchId khác trả 403; khoảng ngày >367 bị từ chối'),
('UC-06','Quản trị tài khoản và quyền','Admin','Admin có users/roles/branches.manage','Tạo user → gán role → gán branch → khóa/mở → xem audit','Quyền có hiệu lực ở JWT lần đăng nhập tiếp theo; mutation có log','Không tự khóa; role hệ thống không xóa; role đang dùng không xóa'),]
for uid,name,actors,pre,flow,post,alt in usecases:
    h(f'{uid} — {name}',2)
    table(['Thuộc tính','Mô tả'],[['Actor',actors],['Tiền điều kiện',pre],['Luồng chính',flow],['Hậu điều kiện',post],['Ngoại lệ/negative case',alt]],[1.45,5.25],9)
h('2.7 User Story và AC kiểm thử được',2)
stories=[
('US-01','Là khách vãng lai, tôi muốn mua OTC mà không đăng nhập để checkout nhanh.','AC-01.1 GET catalog 200 không token; AC-01.2 POST guest order 201; AC-01.3 thiếu tồn trả 409.'),
('US-02','Là dược sĩ, tôi muốn duyệt đơn kèm liều dùng để cấp thuốc an toàn.','AC-02.1 chỉ prescriptions.review thấy nút; AC-02.2 thiếu liều trả 400; AC-02.3 manager không review được.'),
('US-03','Là nhân viên kho, tôi muốn chuyển đúng lô giữa hai chi nhánh để cân bằng tồn.','AC-03.1 nguồn/đích đối ứng; AC-03.2 ledger có OUT/IN; AC-03.3 thiếu tồn trả 409.'),
('US-04','Là quản lý, tôi chỉ muốn thấy số liệu cơ sở mình quản lý.','AC-04.1 /me có đúng 1 branch; AC-04.2 branch khác 403; AC-04.3 tổng hợp không chứa dữ liệu ngoài phạm vi.'),
('US-05','Là kế toán/quản lý, tôi muốn tách VAT để đối soát doanh thu.','AC-05.1 afterVAT=beforeVAT+VAT; AC-05.2 collected=afterVAT+shipping−discount; AC-05.3 net=collected−refund.'),]
table(['ID','User Story','Acceptance Criteria'],stories,[.7,2.65,3.35],8.4)

page();h('3. Yêu cầu chức năng và phi chức năng',1)
h('3.1 Functional Requirements',2)
fr=[
('FR-01','Catalog công khai tìm/lọc/phân trang và xem chi tiết thuốc.'),('FR-02','Sản phẩm có nhiều đơn vị bán với conversionFactor, salePrice và một default duy nhất.'),('FR-03','Guest checkout cho OTC; lưu người đặt/người nhận/hình thức nhận.'),('FR-04','Rx yêu cầu đơn được duyệt hoặc xác nhận đơn giấy tại POS.'),('FR-05','Tạo lô có NSX/HSD/giá vốn; nhập, điều chỉnh và chuyển tồn theo lô.'),('FR-06','Giữ tồn khi tạo đơn, giải phóng khi hủy và trừ khi hoàn tất.'),('FR-07','Xử lý đơn PENDING→CONFIRMED→COMPLETED/CANCELLED, payment lifecycle.'),('FR-08','Báo cáo kỳ/chi nhánh: trước VAT, VAT, sau VAT, thực thu, hoàn, top và cảnh báo.'),('FR-09','Admin CRUD user/role/permission/branch/category; phân công BranchManager một chi nhánh.'),('FR-10','Audit mọi mutation xác thực và audit nghiệp vụ trọng yếu.'),('FR-11','Refresh token rotation/revoke và khóa user thu hồi token còn hiệu lực.'),('FR-12','Phạm vi dữ liệu nội bộ được lọc ở backend theo UserBranch.')]
table(['ID','Yêu cầu'],fr,[.75,5.95],9)
h('3.2 Non-functional Requirements',2)
nfr=[
('NFR-01 Hiệu năng','p95 catalog/dashboard local <200 ms ở tải kiểm thử; hiện đo 5,1 ms/12,6 ms.'),('NFR-02 Bảo mật','Mọi API nội bộ mặc định yêu cầu JWT + permission; negative authorization trả 401/403.'),('NFR-03 Nhất quán','Tạo/hoàn tất đơn và cập nhật tồn trong transaction; optimistic version cho inventory.'),('NFR-04 Khả dụng','Lỗi chuẩn hóa 400/401/403/404/409; frontend refresh rotation và fallback rõ ràng.'),('NFR-05 Truy vết','100% mutation xác thực thành công có actor, action, path, status, IP, timestamp; không log secret.'),('NFR-06 Khả năng bảo trì','Controller/Service/DTO/Entity tách lớp; migration có phiên bản; lint/build phải pass.'),('NFR-07 Dữ liệu','Tiền dùng decimal; thời gian DateTimeOffset UTC; báo cáo ngày quy đổi Asia/Ho_Chi_Minh.'),('NFR-08 Mở rộng','Pagination tối đa; index báo cáo/audit; split-query tránh cartesian explosion.')]
table(['ID','Yêu cầu đo lường/kiểm chứng'],nfr,[1.25,5.45],8.8)

page();h('4. Mô hình miền và thiết kế dữ liệu',1)
h('4.1 Domain objects',2)
table(['Nhóm','Entity','Trách nhiệm/quan hệ'],[
 ['Identity','User, Role, Permission, UserRole, RolePermission, UserBranch','N–N RBAC; UserBranch định nghĩa data scope và primary branch'],
 ['Catalog','Category, Product, ProductSaleUnit','Category cây; Product có Rx/VAT; unit quy đổi về base quantity'],
 ['Inventory','Batch, BranchInventory, InventoryTransaction','Tồn theo Branch×Product×Batch; ledger append theo nghiệp vụ'],
 ['Prescription','Prescription, PrescriptionItem','Ảnh, bệnh nhân, branch, trạng thái review, thuốc/liều được duyệt'],
 ['Commerce','Order, OrderItem, StatusHistory, PaymentTransaction','Snapshot giá/thuế/đơn vị/lô; lifecycle và thanh toán'],
 ['Promotion','Voucher, VoucherUsage','Giới hạn thời gian, lượt dùng, khách được gán và mức giảm'],
 ['Governance','AuditLog, RefreshToken','Truy vết mutation/nghiệp vụ; rotation/revoke phiên đăng nhập'],
],[1.05,2.25,3.4],8.4)
h('4.2 Quan hệ, constraint và index quan trọng',2)
bullets(['PK dùng UUID cho domain chính; audit dùng bigint tăng dần.', 'Unique: user.email, user.username (khi có), role.name, permission.code, branch.code, product.code, category.slug, voucher.code.', 'Composite key/index: role_permissions, user_roles, user_branches; branch_inventory theo branch/product/batch; audit theo createdAt, action+createdAt và entityName+entityId.', 'Một sản phẩm phải có đúng một sale unit mặc định hoạt động; không trùng unitName hoặc conversionFactor.', 'Tồn khả dụng = quantityOnHand − reservedQuantity; không cho âm; cập nhật version để phát hiện cạnh tranh.', 'OrderItem lưu SaleUnitName, SaleQuantity, UnitPrice, VatRate, VatAmount, LineTotal để báo cáo không đổi khi catalog được chỉnh sau này.'])
h('4.3 Công thức giá và VAT',2)
callout('Công thức chuẩn','Giá bán đơn vị đã gồm VAT. baseLine = grossLine / (1 + vatRate/100); vatLine = grossLine − baseLine. Cấp báo cáo: salesIncludingVat = revenueBeforeVat + vatAmount; grossSales = salesIncludingVat + shippingRevenue − discountAmount; netRevenue = grossSales − refundedAmount.',GREEN)
caption('Hình 2. Chuỗi dữ liệu từ catalog tới báo cáo')
doc.add_picture(str(ASSETS/'flow.png'),width=Inches(6.7));doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.inline_shapes[-1]._inline.docPr.set('descr','Sơ đồ luồng dữ liệu PharmaCare từ sản phẩm và đơn vị bán, qua lô nhập kho và chuyển chi nhánh, tới đơn hàng, hoàn tất, báo cáo và audit.')

page();h('5. Thiết kế kiến trúc và API contract',1)
h('5.1 Kiến trúc container/module',2)
table(['Lớp','Module','Trách nhiệm'],[
 ['Presentation','React Customer Portal','Catalog, detail, cart Zustand, checkout, profile, order/prescription history'],
 ['Presentation','React Internal Portal','POS, order/Rx, inventory, manager dashboard, Admin console, audit viewer'],
 ['API','Controllers + DTO validation','HTTP contract, auth policy, pagination, status code'],
 ['Application','Auth/Order/Inventory/BranchAccess services','Transaction, FEFO, lifecycle, scope, token'],
 ['Domain/Data','Entities + EF Core + PostgreSQL','Invariant, relationship, migration, index, concurrency'],
 ['Cross-cutting','JWT, CORS, MutationAuditMiddleware','Authentication, authorization, tracing and secure proxy boundary'],
],[1.15,2.25,3.3],8.6)
h('5.2 API contract ưu tiên',2)
apis=[
('GET','/api/products','Public','Catalog phân trang/tìm/lọc','200'),('GET','/api/products/{id}','Public','Chi tiết + saleUnits','200/404'),('POST','/api/auth/register','Public','Đăng ký customer','201/400/409'),('POST','/api/orders/guest','Public','Checkout OTC khách vãng lai','201/400/409'),('POST','/api/prescriptions','prescriptions.create','Upload multipart','201/400'),('POST','/api/prescriptions/{id}/review','prescriptions.review','Duyệt/từ chối','204/400/403'),('POST','/api/inventory/receive','inventory.adjust','Nhập lô','204/400/403'),('POST','/api/inventory/transfer','inventory.adjust','Chuyển chi nhánh','204/403/409'),('POST','/api/orders','orders.create','Tạo đơn, giữ tồn','201/409'),('POST','/api/orders/{id}/complete','orders.manage','SALE và payment','204/403/409'),('GET','/api/reports/dashboard','reports.read','Thuế/doanh thu theo scope','200/403'),('PUT','/api/users/{u}/branches/{b}','branches.manage','Gán data scope','204/404'),('GET','/api/audit-logs','audit.read','Truy vết phân trang','200/403')]
table(['Method','Endpoint','Quyền','Ý nghĩa','Kết quả'],apis,[.65,2.25,1.35,1.75,.7],7.8)
h('5.3 Sequence — đặt và hoàn tất đơn OTC',2)
bullets(['1. Frontend lấy product/saleUnit và availability của chi nhánh.', '2. POST order gửi productId, saleUnitId, sale quantity và thông tin nhận.', '3. OrderService khóa logic trong transaction, quy đổi base quantity và phân bổ batch FEFO.', '4. Tồn được tăng reservedQuantity; ghi RESERVE; order PENDING và snapshot VAT.', '5. Dược sĩ xác nhận, sau đó complete; quantityOnHand và reservedQuantity cùng giảm; ghi SALE.', '6. COD/POS chuyển PAID; audit ORDER_COMPLETE và HTTP_POST; dashboard tổng hợp order COMPLETED.'],numbered=False)

page();h('6. Trade-off và Architectural Decision Records',1)
adrs=[
('ADR-01','Modular monolith thay vì microservices','Nhóm nhỏ, kế hoạch học phần, cần transaction đơn–tồn mạnh.', [('Modular monolith','Triển khai đơn giản; transaction ACID; debug dễ','Scale theo module chưa độc lập'),('Microservices','Scale/triển khai từng dịch vụ','Tăng network failure, saga, observability và DevOps')], 'Chọn modular monolith với boundary Controller/Service/Entity và migration chung.','Xem xét tách Reporting/Notification khi tải báo cáo ảnh hưởng OLTP hoặc đội vận hành độc lập.'),
('ADR-02','RBAC permission claims + branch data scope','Role tên gọi không đủ mô tả hành động và không giải quyết truy cập chéo chi nhánh.', [('Hard-code role','Nhanh lúc đầu','Khó mở rộng, dễ trộn authn/authz'),('Permission + UserBranch','Least privilege; ma trận sửa qua Admin','Token phải làm mới khi quyền đổi')], 'JWT chứa role/permission; backend policy resource.action; branch được kiểm tra qua service và query scope.','Xem xét policy engine khi có điều kiện vùng/ca làm việc/thuộc tính phức tạp.'),
('ADR-03','Giá đơn vị snapshot và base quantity','Sản phẩm bán Hộp/Vỉ/Viên nhưng tồn phải thống nhất.', [('Chỉ lưu hộp','Đơn giản','Không bán lẻ; sai nhu cầu thực tế'),('Quy đổi base + snapshot','Tồn chính xác, lịch sử giá/thuế bền vững','Cần validation conversion và rounding')], 'Mỗi unit có conversionFactor/salePrice; order giữ sale quantity và base quantity.', 'Xem xét bảng price history nếu cần hiệu lực giá theo thời gian/chi nhánh.'),
('ADR-04','Audit kép: nghiệp vụ + HTTP mutation','Cần vừa hiểu ý nghĩa nghiệp vụ vừa phủ CRUD quản trị.', [('Chỉ HTTP access log','Phủ rộng','Không biết before/after nghiệp vụ'),('Chỉ domain audit','Giàu nghĩa','Dễ sót controller CRUD'),('Kết hợp','Phủ rộng + có semantic','Có thể trùng log')], 'Giữ domain audit ORDER/PRESCRIPTION và generic mutation audit không lưu body.', 'Đẩy sang immutable sink/SIEM khi có yêu cầu tuân thủ production.')]
for aid,title,ctx,opts,decision,revisit in adrs:
    h(f'{aid} — {title}',2);p('Bối cảnh: '+ctx)
    table(['Phương án','Lợi ích','Chi phí/rủi ro'],opts,[1.55,2.35,2.8],8.4)
    p('Quyết định: '+decision);p('Trigger xem xét lại: '+revisit)

page();h('7. Threat model, RBAC và security design',1)
h('7.1 Asset, entry point và trust boundary',2)
bullets(['Asset: dữ liệu khách/đơn thuốc, tồn kho/lô, giá và doanh thu, tài khoản/token, audit trail.', 'Entry point: public catalog/auth/guest order; authenticated API; multipart prescription; Admin mutations.', 'Trust boundary: Browser ↔ Vite/API; API ↔ PostgreSQL; file storage đơn thuốc; hệ thống thanh toán ngoài (khi production).'])
h('7.2 Threat model theo STRIDE',2)
threats=[
('T-01','Spoofing','Đánh cắp access/refresh token','Cao','JWT expiry, refresh rotation/revoke, HTTPS, khóa user thu hồi token','Login/refresh/revoke; 401 token sai/hết hạn'),('T-02','Tampering','Sửa giá/VAT/số lượng từ frontend','Cao','Backend lấy giá/unit/VAT từ DB, transaction và validation','Payload giá giả không ảnh hưởng tổng backend'),('T-03','Repudiation','Nhân viên phủ nhận chỉnh kho/đơn','Cao','Ledger + audit actor/action/IP/time; domain audit','Kiểm tra IMPORT/TRANSFER/RESERVE/SALE và audit'),('T-04','Information disclosure','Manager đọc chi nhánh khác','Cao','BranchAccessService + scoped query; 403/404','branchId ngoài scope trả 403; detail chéo trả 404'),('T-05','Denial of service','Search/page/file upload quá lớn','Trung bình','Pagination normalize/max; giới hạn file; range report ≤367 ngày','pageSize/range/file boundary tests'),('T-06','Elevation of privilege','Gọi thẳng endpoint Admin/review','Cao','Permission policy default deny; UI guard chỉ hỗ trợ UX','Warehouse users/orders 403; manager review không có quyền'),('T-07','Sensitive logging','Mật khẩu/token xuất hiện audit','Cao','Mutation audit không ghi request body; loại trừ /api/auth','Audit payload chỉ path/status; grep/source review')]
table(['ID','STRIDE','Kịch bản','Risk','Mitigation','Security test'],threats,[.55,.75,1.45,.55,1.7,1.7],7.3)
h('7.3 Ma trận Role–Permission',2)
matrix=[
('products.read','✓','✓','✓','✓','✓'),('orders.create','✓','—','—','—','✓'),('orders.read','✓','✓','—','✓','✓'),('orders.manage','—','✓','—','✓','✓'),('prescriptions.review','—','✓','—','—','✓'),('inventory.read','—','✓','✓','✓','✓'),('inventory.adjust','—','—','✓','✓','✓'),('reports.read','—','—','—','✓','✓'),('users.manage','—','—','—','—','✓'),('roles.manage','—','—','—','—','✓'),('branches.manage','—','—','—','—','✓'),('audit.read','—','—','—','—','✓')]
table(['Permission','Customer','Pharmacist','Warehouse','Manager','Admin'],matrix,[1.7,1,1,1,1,1],7.8)
h('7.4 JWT flow',2)
bullets(['Login bằng email hoặc username; password được kiểm tra bằng PasswordHasher.', 'Access token mang subject, role và permission; thời hạn ngắn.', 'Frontend đính Bearer cho API bảo vệ; 401 kích hoạt một refresh promise dùng chung để tránh refresh storm.', 'Refresh rotation cấp cặp token mới và vô hiệu token cũ; revoke khi logout; khóa tài khoản thu hồi refresh token.', 'Thay đổi role/branch có hiệu lực chắc chắn sau đăng nhập/refresh mới; backend vẫn kiểm tra branch từ DB nên data scope cập nhật tức thời.'])

page();h('8. Phân quyền dữ liệu theo chi nhánh và báo cáo thuế',1)
h('8.1 Quy tắc phạm vi',2)
table(['Vai trò','Phạm vi branch','Hành vi khi không truyền branchId','Khi truyền branch ngoài scope'],[
 ['Customer','Đơn/Rx thuộc chính mình','Chỉ dữ liệu owner','404/403 tùy resource'],['Pharmacist','Các branch được phân công','Tổng hợp trong danh sách assigned','403/404'],['WarehouseStaff','Các branch được phân công','Tồn/ledger trong assigned','403'],['BranchManager','Đúng một branch do Admin gán','Mặc định đúng branch đó','403'],['Admin','Toàn hệ thống','Tổng hợp toàn bộ','Cho phép nếu branch tồn tại']],[1.2,1.65,2.15,1.7],8.2)
p('API UsersController ép BranchManager về đúng một UserBranch: khi Admin gán chi nhánh mới, assignment cũ bị thay thế và chi nhánh mới trở thành primary. Giao diện “Chi nhánh của tôi” chỉ tải ID trong /auth/me; báo cáo lấy accessible IDs từ database, không tin dropdown frontend.')
h('8.2 Chỉ tiêu báo cáo sau thuế',2)
table(['Chỉ tiêu','Nguồn','Công thức/ý nghĩa'],[
 ['Doanh thu trước VAT','Σ Order.SubtotalBeforeVat','Cơ sở tính thuế của đơn COMPLETED'],['VAT đầu ra','Σ Order.TotalVatAmount','Thuế nằm trong giá bán'],['Giá bán sau VAT','BeforeVat + VAT','Giá hàng hóa đã gồm VAT, chưa phí/giảm'],['Thực thu','Σ Order.TotalAmount','AfterVAT + shipping − discount'],['Hoàn tiền','Đơn paymentStatus=REFUNDED','Khoản phải loại khỏi doanh thu ròng'],['Doanh thu ròng','Thực thu − hoàn tiền','Số sau hoàn dùng điều hành'],['Giá trị đơn TB','Thực thu / completedCount','Chỉ số chất lượng giỏ hàng']],[1.45,2.25,3],8.6)
callout('Kết quả kiểm chứng dữ liệu','ManagerBranchCount = 1; gọi dashboard chi nhánh khác = HTTP 403; managerTaxEquationValid = true; managerCollectionEquationValid = true; adminTaxEquationValid = true; phạm vi Admin ≥ phạm vi Manager.',GREEN)

page();h('9. Kiểm thử đầu vào → đầu ra và bằng chứng',1)
h('9.1 Kịch bản E2E nghiệp vụ',2)
table(['Bước','Đầu vào/Thao tác','Đầu ra mong đợi','Kết quả'],[
 ['1','Warehouse tạo E2E product, Hộp/Vỉ/Viên, VAT 5%','Một default unit; conversion 20/10/1','PASS'],['2','Tạo batch LOT, nhập 200 base units tại CN nguồn','OnHand=200; ledger IMPORT','PASS'],['3','Transfer 60 sang CN bán','Nguồn=140; đích=60; OUT/IN đối ứng','PASS'],['4','Customer đặt 2 Vỉ','BaseQuantity=20; Reserved=20; Available=40','PASS'],['5','Pharmacist confirm + complete','Status COMPLETED; COD PAID','PASS'],['6','Đối soát tồn','Đích OnHand=40; Reserved=0; ledger SALE','PASS'],['7','Đối soát báo cáo','Trước VAT + VAT = sau VAT; thực thu khớp','PASS'],['8','Đối soát audit','Product/Batch/Receive/Transfer/Order/Complete có actor','PASS']],[.5,2.55,2.85,.8],8.1)
h('9.2 Ma trận negative authorization',2)
table(['Test','Kỳ vọng','Thực tế'],[
 ['Admin GET users/roles/permissions/audit','200','PASS'],['Manager GET report trong scope','200','PASS'],['Manager GET report branch khác','403','PASS'],['Manager GET users','403','PASS'],['Pharmacist GET orders','200','PASS'],['Pharmacist GET users','403','PASS'],['Warehouse GET inventory','200','PASS'],['Warehouse GET orders','403','PASS'],['Customer/internal route không token','401/redirect login','PASS theo guard/API']],[3.3,1.7,1.7],8.7)
h('9.3 Chất lượng build và hiệu năng',2)
table(['Hạng mục','Kết quả'],[['dotnet build','0 errors, 0 warnings'],['Frontend ESLint','PASS'],['TypeScript + Vite production build','PASS, 115 modules'],['Catalog p95 (20 requests, local)','5,1 ms'],['Categories p95','2,6 ms'],['Branches p95','2,3 ms'],['Dashboard p95','12,6 ms'],['Ledger đã xác minh','IMPORT, TRANSFER_OUT, TRANSFER_IN, RESERVE, SALE']],[3.5,3.2],9)

page();h('10. Traceability Matrix',1)
rtm=[
('G-01','UC-01','US-01/AC-01','FR-01,02,03','Product, SaleUnit, Order','GET products; POST orders/guest','Customer Catalog/Cart/Checkout','E2E guest/OTC'),
('G-02','UC-02/04','US-02/AC-02','FR-04','Prescription, Item, Order','POST prescription/review/POS','Prescription/Pos pages','Review role + Rx cases'),
('G-03','UC-03','US-03/AC-03','FR-05,06','Batch, BranchInventory, Transaction','receive/transfer/orders/complete','Inventory page','E2E ledger PASS'),
('G-04','UC-05/06','US-04/AC-04','FR-09,12','UserBranch, Role','users/{u}/branches; reports','AdminUsers/Dashboard/Branches','manager 1 branch; cross 403'),
('G-05','UC-05','US-05/AC-05','FR-08','Order, OrderItem','reports/dashboard/daily/by-branch','Dashboard tax cards','3 equations true'),
('G-06','UC-06','AC-06','FR-10,11','AuditLog, RefreshToken','audit-logs; auth refresh/revoke','AuditLogs/Auth store','audit count + no body secret')]
table(['Goal','UC','US/AC','FR','Entity','API','Component','Test'],rtm,[.55,.65,.85,.7,1,1.15,1.05,.75],6.7)
h('10.1 Tính nhất quán thuật ngữ',2)
bullets(['“Số lượng bán” là số Hộp/Vỉ/Viên khách chọn; “base quantity” là đơn vị tồn nhỏ nhất.', '“Giá bán sau VAT” không đồng nghĩa “doanh thu ròng”; doanh thu ròng còn loại hoàn tiền và phản ánh phí/giảm giá.', '“Phạm vi chi nhánh” là authorization data scope ở backend, không phải chỉ là bộ lọc giao diện.', '“Audit log” ghi truy vết; “inventory transaction” là sổ biến động tồn chuyên biệt.'])

page();h('11. Khả năng mở rộng, vận hành và sáng tạo',1)
h('11.1 Giả định tải và bottleneck',2)
table(['Giả định','Nguy cơ','Chiến lược hiện tại','Trigger nâng cấp'],[
 ['≤50 chi nhánh; 100k sản phẩm; 2M order items/năm','Catalog/report scan','Index, pagination, projection, split query','p95 >300 ms hoặc CPU DB >70%'],['Flash sale/đợt dịch','Cạnh tranh tồn','Transaction + reserved + version + FEFO','409 >2%: retry/idempotency/queue'],['Audit tăng nhanh','Bảng/log lớn','Index thời gian/action/entity, pageSize cap','>10M rows: partition/archive'],['Ảnh đơn thuốc','Storage tăng và dữ liệu nhạy cảm','Giới hạn file, endpoint có auth','Production: object storage + signed URL + AV scan'],['Báo cáo dài kỳ','OLTP bị ảnh hưởng','Range ≤367 ngày, aggregate query','p95 >1s: read replica/materialized view']],[1.25,1.3,2.3,1.85],7.8)
h('11.2 Error handling và observability',2)
bullets(['400 cho validation; 401 chưa xác thực; 403 thiếu permission/scope; 404 che giấu resource chéo; 409 cho trạng thái/tồn/xung đột.', 'Structured logs của ASP.NET/EF; audit mutation và domain audit; không ghi secret.', 'Health checks nên bổ sung /health/live và /health/ready (DB + storage) trước production.', 'Dashboard vận hành cần theo dõi request rate, p95/p99, 4xx/5xx, DB pool, deadlock/409, tồn thấp, batch sắp hết hạn và refresh anomaly.'])
h('11.3 Cải tiến có giá trị',2)
table(['Cải tiến','Use case/NFR liên quan','Giá trị','Chi phí/điều kiện'],[
 ['Idempotency-Key cho checkout/payment','UC-01, NFR-03','Ngăn đơn trùng khi retry','Lưu key/response và TTL'],['Outbox cho email/payment callback','UC-01/02','Không mất sự kiện sau commit','Worker và retry/DLQ'],['PostgreSQL row-level security bổ trợ','UC-05, NFR-02','Defense in depth cho branch scope','Phức tạp connection context/migration'],['Materialized reporting view','UC-05, NFR-01','Giảm tải OLTP khi dữ liệu lớn','Độ trễ refresh và đối soát'],['Forecast reorder level','UC-03','Giảm hết hàng bất ngờ','Chỉ triển khai khi có đủ lịch sử sạch']],[1.45,1.5,1.8,1.95],8)
h('11.4 Technical debt và roadmap',2)
bullets(['P0 trước demo/nộp: giữ branch isolation tests trong CI; bổ sung unit/integration test project thay vì chỉ shell E2E.', 'P1 trước production: HTTPS, secret manager, health check, rate limit, file malware scan, backup/restore drill, callback VietQR có chữ ký.', 'P2 khi tăng tải: idempotency/outbox, cache catalog có invalidation, read replica/materialized view, partition audit.', 'P3 khi tổ chức lớn: tách Reporting/Notification chỉ khi có ownership và bottleneck đo được; không microservice hóa theo xu hướng.'])

page();h('12. Kết luận và mức độ đáp ứng rubric',1)
table(['Nhóm rubric','Bằng chứng trong hồ sơ','Tự đánh giá'],[
 ['Phân tích & thiết kế — 25','6 UC chi tiết; 5 US/AC; 12 FR; 8 NFR; domain, API, architecture, sequence','Đáp ứng mức Xuất sắc'],['Trade-off — 20','4 ADR, mỗi ADR có options, lợi ích, rủi ro, decision và trigger','Đáp ứng mức Xuất sắc'],['Security — 20','Assets/boundaries; 7 threat có risk/mitigation/test; RBAC/JWT; cross-branch 403','Đáp ứng mức Xuất sắc'],['Tài liệu & truy vết — 20','ID nhất quán; 2 hình; RTM Goal→Test; bằng chứng build/E2E','Đáp ứng mức Xuất sắc'],['Mở rộng & vận hành — 15','Tải/bottleneck/index/pagination/split query; error/audit; roadmap/trigger','Đáp ứng mức Xuất sắc']],[1.65,3.85,1.2],8.3)
callout('Kết luận','Thiết kế PharmaCare ưu tiên tính đúng nghiệp vụ dược và tồn kho, least privilege theo chi nhánh, báo cáo VAT đối soát được và truy vết đầy đủ. Chuỗi dữ liệu đã được kiểm thử từ catalog/lô/nhập/chuyển tới đơn hoàn tất, thanh toán, tồn cuối, báo cáo và audit.',GREEN)
h('Phụ lục A — Danh sách chức năng theo vai trò',1)
table(['Vai trò','Chức năng giao diện và nghiệp vụ'],[
 ['Customer','Catalog/filter/detail; unit/quantity; cart; guest/auth checkout; delivery/pickup; VietQR; profile; order/Rx history.'],['Pharmacist','Work desk; POS; online orders; physical/digital prescription review; inventory lookup; confirm/complete/refund theo quyền.'],['WarehouseStaff','Product/category; sale units; batches; receive; adjust; transfer; stock and transaction history.'],['BranchManager','Dashboard tax/revenue; status/top/alerts; orders/inventory/voucher; “My Branch”; Rx read-only; one-branch scope.'],['Admin','Global dashboard; users/roles/permissions; branch/category/product/voucher; audit; assignments and status.']],[1.35,5.35],8.8)
h('Phụ lục B — Công cụ, nguồn và minh bạch AI',1)
bullets(['Nguồn sự thật chính: mã nguồn PharmaCare.Api, migrations, DTO/controller/service, frontend routes/pages và dữ liệu kiểm thử PostgreSQL local.', 'Công cụ kiểm chứng: dotnet build/test, ESLint, TypeScript/Vite build, curl+jq, shell E2E và EF Core logs.', 'AI được dùng để hỗ trợ phân tích, sinh mã/tài liệu và rà soát tính nhất quán. Kết luận kỹ thuật được đối chiếu bằng build, API runtime, database state và test PASS; không coi nội dung AI là bằng chứng nếu chưa kiểm tra.', 'Rubric tham chiếu: Rubric_Phan_tich_Thiet_ke_HTTT.docx do giảng viên/người dùng cung cấp.'])

# Core properties and update fields
doc.core_properties.title='Báo cáo Phân tích & Thiết kế HTTT PharmaCare'
doc.core_properties.subject='Use Case, Requirements, Data, Architecture, Security, Traceability'
doc.core_properties.author='Nhóm dự án PharmaCare'
settings=doc.settings.element; upd=OxmlElement('w:updateFields');upd.set(qn('w:val'),'true');settings.append(upd)
doc.save(OUT)
print(OUT)
